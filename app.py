import streamlit as st
import openai
import google.generativeai as genai
import gspread
import datetime
import json
import os
import time
from docx import Document
from io import BytesIO

# -----------------------------------------------------------
# 0. [설정] 가격표 및 스타일
# -----------------------------------------------------------
PRICES = {
    "chat_step0": 10,      # 토론 코멘트
    "var_confirm": 25,     # 변인 확정 (구매)
    "method_confirm": 30,  # 방법 확정 (구매)
    "search": 30,          # 검색
    "draft": 100,          # 논문 작성
    "ref": 30,             # 참고문헌
    "side_chat": 5         # 조교 질문
}

st.set_page_config(page_title="MJP Research Lab", layout="wide")

st.markdown("""
<style>
    div.stButton > button:first-child { background-color: #2c3e50; color: white; border-radius: 6px; border: none; font-weight: 600;}
    div.stButton > button:first-child:hover { background-color: #1a252f; }
    .energy-box { padding: 12px 20px; background-color: #f8f9fa; border-left: 5px solid #2c3e50; border-radius: 4px; display: flex; align-items: center; gap: 15px; margin-bottom: 25px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
    .energy-val { font-size: 22px; font-weight: bold; color: #2c3e50; font-family: monospace; }
    .confirm-box { padding: 15px; border: 2px solid #e74c3c; background-color: #fdedec; border-radius: 8px; margin-top: 10px; margin-bottom: 10px; text-align: center; }
</style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------
# 1. 시스템 함수 (DB, Log)
# -----------------------------------------------------------
@st.cache_resource
def get_google_sheet_connection():
    try:
        if "gcp_service_account" not in st.secrets: return None
        gc = gspread.service_account_from_dict(st.secrets["gcp_service_account"])
        return gc.open("MJP 연구실 관리대장")
    except: return None

def fetch_users_from_sheet():
    sh = get_google_sheet_connection()
    if not sh: return {"zenova90": "0931285asd*"}
    try:
        ws = sh.worksheet("Users")
        records = ws.get_all_values()
        user_dict = {}
        for row in records[1:]:
            if len(row) >= 3: user_dict[row[1]] = row[2]
        user_dict["zenova90"] = "0931285asd*"
        return user_dict
    except: return {"zenova90": "0931285asd*"}

def register_user_to_sheet(new_id, new_pw):
    sh = get_google_sheet_connection()
    if not sh: return False, "DB 연동 오류"
    current = fetch_users_from_sheet()
    if new_id in current: return False, "❌ 이미 존재하는 ID입니다."
    try:
        ws = sh.worksheet("Users")
        ws.append_row([datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), new_id, new_pw])
        return True, "✅ 가입 완료!"
    except: return False, "가입 실패"

def log_to_sheet(username, action, content):
    sh = get_google_sheet_connection()
    if not sh: return
    try:
        ws = sh.worksheet("Logs")
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ws.append_row([datetime.datetime.now().strftime("%Y-%m-%d"), ts, username, action, content])
    except: pass

def fetch_logs_by_date(username, target_date_str):
    sh = get_google_sheet_connection()
    if not sh: return []
    try:
        ws = sh.worksheet("Logs")
        rows = ws.get_all_values()
        filtered = []
        for row in rows[1:]:
            if len(row) >= 5 and row[0] == target_date_str and row[2] == username:
                filtered.append({"time": row[1], "action": row[3], "content": row[4]})
        return sorted(filtered, key=lambda x: x['time'], reverse=True)
    except: return []

# -----------------------------------------------------------
# 2. 워드 생성 및 유틸
# -----------------------------------------------------------
def create_word_report(username, date_str, logs):
    doc = Document()
    doc.add_heading(f'{username}님의 연구 일지', 0)
    doc.add_paragraph(f'날짜: {date_str}')
    if not logs: doc.add_paragraph("기록 없음")
    else:
        for log in logs:
            doc.add_heading(f"[{log['time']}] {log['action']}", level=2)
            doc.add_paragraph(log['content'])
            doc.add_paragraph("-" * 30)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def check_and_deduct(cost):
    if st.session_state['user_energy'] >= cost:
        st.session_state['user_energy'] -= cost
        return True
    st.error(f"⚠️ 에너지가 부족합니다. (필요: {cost})"); return False

# -----------------------------------------------------------
# 3. AI 함수
# -----------------------------------------------------------
openai.api_key = st.secrets.get("OPENAI_API_KEY", "")
genai.configure(api_key=st.secrets.get("GEMINI_API_KEY", ""))

def chat_with_context(prompt, context_data, stage_name):
    try:
        sys_msg = f"당신은 심리학 연구 조교 '다온'입니다.\n단계: {stage_name}\n[화면 내용]\n{context_data}"
        res = openai.chat.completions.create(model="gpt-4o-mini", messages=[{"role":"system","content":sys_msg},{"role":"user","content":prompt}])
        return res.choices[0].message.content
    except: return "AI 오류"

def get_ai_options_4(prompt): # 4개 제안으로 변경
    try:
        res = openai.chat.completions.create(model="gpt-4o-mini", messages=[{"role":"user","content":prompt}])
        # 줄바꿈이나 |||로 분리 시도
        content = res.choices[0].message.content
        if "|||" in content: return [opt.strip() for opt in content.split("|||") if opt.strip()]
        else: return [opt.strip() for opt in content.split("\n") if opt.strip()][:4]
    except: return ["제안 실패"]

def search_literature(topic, vars_text):
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        return model.generate_content(f"주제: {topic}, 변인: {vars_text}. 선행연구 3개 검색 요약.").text
    except: return "검색 오류"

# -----------------------------------------------------------
# 4. 세션 초기화
# -----------------------------------------------------------
if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
if 'username' not in st.session_state: st.session_state['username'] = ""
if 'user_energy' not in st.session_state: st.session_state['user_energy'] = 500

if 'research_context' not in st.session_state: st.session_state['research_context'] = {}
for k in ['topic', 'variables_options', 'variables', 'method_options', 'method', 'references']:
    if k not in st.session_state['research_context']:
        st.session_state['research_context'][k] = [] if 'options' in k else ""

if 'paper_sections' not in st.session_state:
    st.session_state['paper_sections'] = {"서론": "", "이론적 배경": "", "연구 방법": "", "결과": "", "논의": ""}

for k in ["chat_0", "chat_1", "chat_2", "chat_3", "chat_4", "chat_5"]:
    if k not in st.session_state: st.session_state[k] = []

# 확인 팝업 상태 관리
if 'confirm_state' not in st.session_state: st.session_state['confirm_state'] = {"type": None, "data": None}

# -----------------------------------------------------------
# 5. 메인 앱
# -----------------------------------------------------------
def login_page():
    st.title("🔐 MJP Research Lab")
    t1, t2 = st.tabs(["로그인", "회원가입 (자율)"])
    with t1:
        with st.form("login"):
            uid = st.text_input("아이디")
            upw = st.text_input("비밀번호", type="password")
            if st.form_submit_button("로그인"):
                users = fetch_users_from_sheet()
                if uid in users and users[uid] == upw:
                    st.session_state['logged_in'] = True; st.session_state['username'] = uid
                    log_to_sheet(uid, "로그인", "성공")
                    st.rerun()
                else: st.error("정보 불일치")
    with t2:
        with st.form("signup"):
            nid = st.text_input("희망 아이디")
            npw = st.text_input("희망 비밀번호", type="password")
            if st.form_submit_button("가입하기"):
                suc, msg = register_user_to_sheet(nid, npw)
                if suc: st.success(msg)
                else: st.error(msg)

def render_right_chat(key_suffix, context_data, stage_name):
    st.markdown(f"###### 💬 AI 조교 ({stage_name})")
    cost = PRICES["side_chat"]
    
    chat_key = f"chat_{key_suffix}"
    for msg in st.session_state[chat_key]:
        with st.chat_message(msg["role"]): st.markdown(msg["content"])
        
    if p := st.chat_input(f"질문하기 (비용: {cost}E)", key=f"in_{key_suffix}"):
        if check_and_deduct(cost):
            st.session_state[chat_key].append({"role":"user", "content":p})
            log_to_sheet(st.session_state['username'], f"질문({stage_name})", p)
            with st.chat_message("user"): st.markdown(p)
            with st.spinner("..."):
                ans = chat_with_context(p, context_data, stage_name)
                st.session_state[chat_key].append({"role":"assistant", "content":ans})
                log_to_sheet(st.session_state['username'], f"답변({stage_name})", ans)
                st.rerun()

def main_app():
    user = st.session_state['username']
    
    # [사이드바]
    with st.sidebar:
        st.header(f"👤 {user}님")
        st.markdown("---")
        st.subheader("📅 연구 기록")
        search_date = st.date_input("날짜")
        if st.button("기록 보기"):
            logs = fetch_logs_by_date(user, search_date.strftime("%Y-%m-%d"))
            st.session_state['fetched_logs'] = logs; st.session_state['fetched_date'] = search_date.strftime("%Y-%m-%d")
            if not logs: st.info("기록 없음")
            
        if 'fetched_logs' in st.session_state and st.session_state['fetched_logs']:
            docx = create_word_report(user, st.session_state['fetched_date'], st.session_state['fetched_logs'])
            st.download_button("📄 워드 다운로드", docx, f"Log_{st.session_state['fetched_date']}.docx")

        # [NEW] 수동 저장 버튼
        st.markdown("---")
        if st.button("💾 오늘의 연구 기록 저장"):
            summary = f"Topic: {st.session_state['research_context']['topic']}\nVars: {st.session_state['research_context']['variables']}\nMethod: {st.session_state['research_context']['method']}"
            log_to_sheet(user, "수동저장", summary)
            st.success("현재 상태가 캘린더에 저장되었습니다.")

        if user == "zenova90":
            st.markdown("---")
            st.error("🔒 관리자")
            st.link_button("📂 구글 시트 열기", "https://docs.google.com/spreadsheets")

        st.markdown("---")
        with st.expander("⚡ 충전소"):
            code = st.text_input("쿠폰")
            if st.button("충전"):
                if code == "TEST-1000":
                    st.session_state['user_energy'] += 1000
                    log_to_sheet(user, "충전", "1000E"); st.success("완료")
        
        if st.button("로그아웃"): st.session_state['logged_in'] = False; st.rerun()

    # [메인]
    st.title("🎓 MJP Research Lab")
    st.markdown(f"<div class='energy-box'><span>⚡ Energy:</span><span class='energy-val'>{st.session_state['user_energy']}</span></div>", unsafe_allow_html=True)
    
    tabs = st.tabs(["💡 0. 토론", "1. 변인", "2. 방법", "3. 검색", "4. 작성", "5. 참고", "📜 로그"])

    # [Tab 0: 토론]
    with tabs[0]:
        st.header("💡 Brainstorming")
        cost = PRICES["chat_step0"]
        render_right_chat("0", "초기 단계", "토론") 

    # [Tab 1: 변인]
    with tabs[1]:
        cL, cR = st.columns([6, 4])
        with cL:
            st.subheader("Variables")
            
            # 1. 제안 생성 (무료)
            topic = st.text_input("연구 주제", value=st.session_state['research_context']['topic'])
            if st.button("🤖 4가지 안 제안받기 (무료)", key="ai_v_free"):
                with st.spinner("생성 중..."):
                    opts = get_ai_options_4(f"주제 '{topic}'에 적합한 변인 구조 4가지를 제안해줘. 각 안은 |||로 구분해줘.")
                    st.session_state['research_context']['variables_options'] = opts
                    st.rerun()

            # 2. 선택 및 확정 (유료)
            if st.session_state['research_context']['variables_options']:
                choice = st.radio("마음에 드는 안을 선택하세요:", st.session_state['research_context']['variables_options'])
                
                # 확정 시도 버튼
                if st.button("선택한 안으로 적용하기"):
                    st.session_state['confirm_state'] = {"type": "var", "data": choice}
                    st.rerun()

            # 3. 재확인 팝업 (가짜 팝업 구현)
            if st.session_state['confirm_state']['type'] == "var":
                st.markdown(f"""
                <div class='confirm_box'>
                    <h4>💰 확인 필요</h4>
                    <p>변인을 확정하면 <b>{PRICES['var_confirm']} 에너지</b>가 차감됩니다.</p>
                    <p>선택한 내용: {st.session_state['confirm_state']['data'][:20]}...</p>
                </div>
                """, unsafe_allow_html=True)
                
                col_y, col_n = st.columns(2)
                if col_y.button("✅ 네, 결제하고 적용합니다"):
                    if check_and_deduct(PRICES['var_confirm']):
                        st.session_state['research_context']['variables'] = st.session_state['confirm_state']['data']
                        log_to_sheet(user, "변인확정", st.session_state['confirm_state']['data'])
                        st.session_state['confirm_state'] = {"type": None, "data": None} # 초기화
                        st.success("적용 완료!")
                        st.rerun()
                if col_n.button("❌ 아니오, 취소합니다"):
                    st.session_state['confirm_state'] = {"type": None, "data": None}
                    st.rerun()

            st.text_area("최종 확정된 변인", value=st.session_state['research_context']['variables'], height=150)

        with cR:
            render_right_chat("1", f"주제:{topic}\n현재변인:{st.session_state['research_context']['variables']}", "변인")

    # [Tab 2: 방법]
    with tabs[2]:
        cL, cR = st.columns([6, 4])
        with cL:
            st.subheader("Methodology")
            
            # 제안 (무료)
            if st.button("🤖 4가지 방법론 제안받기 (무료)", key="ai_m_free"):
                with st.spinner("생성 중..."):
                    opts = get_ai_options_4(f"변인 '{st.session_state['research_context']['variables']}'에 맞는 연구방법 4가지 제안 (|||로 구분)")
                    st.session_state['research_context']['method_options'] = opts
                    st.rerun()
            
            if st.session_state['research_context']['method_options']:
                choice_m = st.radio("방법론 선택:", st.session_state['research_context']['method_options'])
                
                if st.button("선택한 방법론 적용"):
                    st.session_state['confirm_state'] = {"type": "method", "data": choice_m}
                    st.rerun()

            # 재확인 팝업
            if st.session_state['confirm_state']['type'] == "method":
                st.markdown(f"""
                <div class='confirm_box'>
                    <h4>💰 확인 필요</h4>
                    <p>방법론을 확정하면 <b>{PRICES['method_confirm']} 에너지</b>가 차감됩니다.</p>
                </div>""", unsafe_allow_html=True)
                
                col_y, col_n = st.columns(2)
                if col_y.button("✅ 결제 및 적용"):
                    if check_and_deduct(PRICES['method_confirm']):
                        st.session_state['research_context']['method'] = st.session_state['confirm_state']['data']
                        log_to_sheet(user, "방법확정", st.session_state['confirm_state']['data'])
                        st.session_state['confirm_state'] = {"type": None, "data": None}
                        st.success("적용 완료!")
                        st.rerun()
                if col_n.button("❌ 취소"):
                    st.session_state['confirm_state'] = {"type": None, "data": None}
                    st.rerun()

            st.text_area("최종 방법", value=st.session_state['research_context']['method'])

        with cR:
            render_right_chat("2", f"방법:{st.session_state['research_context']['method']}", "방법")

    # [Tab 3: 검색]
    with tabs[3]:
        cL, cR = st.columns([6, 4])
        with cL:
            st.subheader("Search")
            cost = PRICES['search']
            if st.button(f"🚀 Gemini 검색 ({cost}E)", key="s_g"):
                if check_and_deduct(cost):
                    res = search_literature(st.session_state['research_context']['topic'], st.session_state['research_context']['variables'])
                    st.session_state['research_context']['references'] = res
                    log_to_sheet(user, "검색", res)
                    st.rerun()
            st.text_area("결과", value=st.session_state['research_context']['references'])
        with cR: render_right_chat("3", st.session_state['research_context']['references'], "검색")

    # [Tab 4: 작성]
    with tabs[4]:
        cL, cR = st.columns([6, 4])
        with cL:
            st.subheader("Drafting")
            sec = st.selectbox("챕터", list(st.session_state['paper_sections'].keys()))
            cost = PRICES['draft']
            
            # 재확인 로직 (작성은 비싸니까!)
            if st.button(f"🤖 AI 초안 작성 요청"):
                st.session_state['confirm_state'] = {"type": "draft", "data": sec}
                st.rerun()
                
            if st.session_state['confirm_state']['type'] == "draft":
                st.markdown(f"<div class='confirm_box'><h4>💰 {cost} 에너지 차감</h4><p>'{sec}' 챕터를 작성하시겠습니까?</p></div>", unsafe_allow_html=True)
                cy, cn = st.columns(2)
                if cy.button("✅ 작성 시작"):
                    if check_and_deduct(cost):
                        st.session_state['confirm_state'] = {"type": None, "data": None}
                        with st.spinner("작성 중..."):
                            draft = chat_with_context(f"'{sec}' 챕터 작성해줘", str(st.session_state['research_context']), "작성")
                            st.session_state['paper_sections'][sec] = draft
                            log_to_sheet(user, f"작성({sec})", draft)
                            st.rerun()
                if cn.button("❌ 취소"):
                    st.session_state['confirm_state'] = {"type": None, "data": None}; st.rerun()
            
            cur = st.text_area("에디터", value=st.session_state['paper_sections'][sec])
            if st.button("저장"): st.session_state['paper_sections'][sec]=cur; log_to_sheet(user, f"수정({sec})", cur); st.success("저장됨")

        with cR: render_right_chat("4", f"챕터:{sec}\n{st.session_state['paper_sections'][sec]}", "작성")

    # [Tab 5: 참고문헌]
    with tabs[5]:
        cL, cR = st.columns([6, 4])
        with cL:
            cost = PRICES['ref']
            if st.button(f"✨ APA 변환 ({cost}E)"):
                if check_and_deduct(cost):
                    res = chat_with_context("APA 변환해줘", st.session_state['research_context']['references'], "참고문헌")
                    st.markdown(res)
        with cR: render_right_chat("5", st.session_state['research_context']['references'], "참고")

    # [Tab 6: 로그]
    with tabs[6]:
        st.header("Today's Log")
        today = datetime.datetime.now().strftime("%Y-%m-%d")
        logs = st.session_state.get('fetched_logs', fetch_logs_by_date(user, today))
        for log in logs:
            st.markdown(f"<div class='log-entry'><b>{log['time']}</b> [{log['action']}]<br>{log['content'][:100]}...</div>", unsafe_allow_html=True)

if st.session_state['logged_in']: main_app()
else: login_page()
